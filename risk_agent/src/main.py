from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import os
from doc_parser import process_documents
from pathlib import Path
import asyncio
from concurrent.futures import ThreadPoolExecutor
import json
from typing import List

from agents import chat_with_rag
from rag_ingest import build_vector_store

from docx import Document
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

app = FastAPI()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allows all methods
    allow_headers=["*"],  # Allows all headers
)

UPLOAD_DIR = "uploaded_docs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
TEMP_DIR = "temp_reports"
os.makedirs(TEMP_DIR, exist_ok=True)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")


executor = ThreadPoolExecutor(max_workers=4)

async def delete_file_after_delay(path: str, delay: int):
    await asyncio.sleep(delay)
    try:
        os.remove(path)
        print(f"Successfully deleted temporary file: {path}")
    except OSError as e:
        print(f"Error deleting file {path}: {e}")

@app.get("/")
def root():
    return FileResponse('static/index.html')


@app.post("/upload")
async def upload_documents(files: List[UploadFile] = File(...)):
    if len(files) != 3:
        raise HTTPException(status_code=400, detail="Exactly three files are required.")

    saved_files = []
    for uploaded_file in files:
        file_path = Path(UPLOAD_DIR) / uploaded_file.filename
        with open(file_path, "wb") as buffer:
            buffer.write(await uploaded_file.read())
        saved_files.append(file_path.resolve())

    scope_path, requirements_path, risks_path = map(str, saved_files)

    try:
        md_files = await asyncio.get_running_loop().run_in_executor(
            executor,
            process_documents,
            scope_path,
            requirements_path,
            risks_path,
        )

        collection = build_vector_store()
        app.state.collection = collection

        return {
            "message": "Documents processed and indexed successfully",
            "markdown_files": md_files,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/risk-register")
async def generate_risk_register():
    if not hasattr(app.state, "collection"):
        raise HTTPException(
            status_code=400,
            detail="Documents not ingested yet. Upload documents first.",
        )

    question = (
        "Generate a new risk register for the project "
        "based on current scope and requirements."
    )

    try:
        raw_output = await chat_with_rag(
            app.state.collection,
            question,
        )

        try:
            parsed_json = json.loads(raw_output)
            app.state.risk_register_data = parsed_json # Cache the data
        except json.JSONDecodeError:
            raise HTTPException(
                status_code=500,
                detail=f"Model returned invalid JSON:\n{raw_output}",
            )

        return parsed_json

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate risk register: {e}",
        )

@app.get("/download-report")
async def download_report(format: str, background_tasks: BackgroundTasks):
    if not hasattr(app.state, "risk_register_data"):
        raise HTTPException(
            status_code=400,
            detail="Risk register not generated yet.",
        )
    
    data = app.state.risk_register_data
    headers = list(data[0].keys())
    table_data = [headers] + [[str(item.get(h, '')) for h in headers] for item in data]

    if format == "docx":
        document = Document()
        document.add_heading('Risk Register', 0)
        table = document.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        for row_data in data:
            row_cells = table.add_row().cells
            for i, header in enumerate(headers):
                row_cells[i].text = str(row_data.get(header, ''))

        file_path = os.path.join(TEMP_DIR, "risk_register.docx")
        document.save(file_path)
        background_tasks.add_task(delete_file_after_delay, file_path, 300)
        return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename='risk_register.docx')

    elif format == "pdf":
        file_path = os.path.join(TEMP_DIR, "risk_register.pdf")
        doc = SimpleDocTemplate(file_path, pagesize=landscape(letter))
        
        # Prepare data for the table
        styles = getSampleStyleSheet()
        styleN = styles['BodyText']
        styleN.alignment = 1  # Center alignment
        
        # Wrap data in Paragraphs
        wrapped_data = []
        for row in table_data:
            wrapped_row = [Paragraph(cell, styleN) for cell in row]
            wrapped_data.append(wrapped_row)
        
        table = Table(wrapped_data)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('WORDWRAP', (0, 0), (-1, -1)),
        ])
        table.setStyle(style)
        
        elements = [table]
        doc.build(elements)
        background_tasks.add_task(delete_file_after_delay, file_path, 300)
        return FileResponse(file_path, media_type='application/pdf', filename='risk_register.pdf')

    else:
        raise HTTPException(status_code=400, detail="Invalid format specified. Use 'pdf' or 'docx'.")
